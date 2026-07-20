from __future__ import annotations
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from ..document.model import DocumentModel, DocumentSection, Paragraph, Table, KeyValueTable, PageBreak


class WordExporter:
    """Renders a presentation-neutral DocumentModel to a clean DOCX."""

    def export(self, model: DocumentModel, output_path: str | Path) -> Path:
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        self._configure_document(doc)
        self._add_cover(doc, model)
        for section in model.sections:
            self._add_section(doc, section)
        self._add_footer(doc, model)
        doc.save(path)
        return path

    def _configure_document(self, doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.8)

        styles = doc.styles
        normal = styles['Normal']
        normal.font.name = 'Arial'
        normal.font.size = Pt(9)
        normal.paragraph_format.space_after = Pt(4)
        normal.paragraph_format.line_spacing = 1.05
        for name, size, before, after in [
            ('Title', 24, 0, 18), ('Subtitle', 11, 0, 12),
            ('Heading 1', 16, 16, 8), ('Heading 2', 12, 12, 6), ('Heading 3', 10, 10, 4),
        ]:
            style = styles[name]
            style.font.name = 'Arial'
            style.font.size = Pt(size)
            style.font.bold = name.startswith('Heading') or name == 'Title'
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

    def _add_cover(self, doc: Document, model: DocumentModel) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(72)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(model.title)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(24)

        if model.subtitle:
            p = doc.add_paragraph(model.subtitle)
            p.style = doc.styles['Subtitle']

        doc.add_paragraph('')
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        details = [
            ('Project', model.project_name or '-'),
            ('Source', model.source_name or '-'),
            ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M')),
            ('Generator', f'{model.generated_by} {model.version}'.strip()),
        ]
        for key, value in details:
            cells = table.add_row().cells
            cells[0].width = Cm(3.2)
            cells[1].width = Cm(11.5)
            cells[0].text = key
            cells[1].text = value
            cells[0].paragraphs[0].runs[0].bold = True
        self._style_table(table, header=False)
        doc.add_page_break()

    def _add_section(self, doc: Document, section: DocumentSection) -> None:
        level = max(1, min(section.level, 3))
        doc.add_heading(section.title, level=level)
        for element in section.elements:
            if isinstance(element, Paragraph):
                self._add_paragraph(doc, element)
            elif isinstance(element, Table):
                self._add_table(doc, element)
            elif isinstance(element, KeyValueTable):
                self._add_key_value_table(doc, element)
            elif isinstance(element, PageBreak):
                doc.add_page_break()
        for child in section.sections:
            self._add_section(doc, child)

    def _add_paragraph(self, doc: Document, element: Paragraph) -> None:
        p = doc.add_paragraph(element.text)
        p.paragraph_format.keep_with_next = element.keep_with_next
        if element.style == 'note':
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.italic = True
        elif element.style == 'compact':
            p.paragraph_format.space_after = Pt(1)

    def _add_table(self, doc: Document, element: Table) -> None:
        if element.title:
            p = doc.add_paragraph(element.title)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(5)
            p.runs[0].bold = True
        table = doc.add_table(rows=1, cols=len(element.columns))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        hdr = table.rows[0].cells
        for i, label in enumerate(element.columns):
            hdr[i].text = label
        if element.repeat_header:
            self._repeat_table_header(table.rows[0])
        for row in element.rows:
            cells = table.add_row().cells
            for i, value in enumerate(row.cells[:len(cells)]):
                cells[i].text = str(value)
        if element.widths:
            for row in table.rows:
                for i, width in enumerate(element.widths[:len(row.cells)]):
                    row.cells[i].width = Cm(width * 2.15)
        self._style_table(table, header=True)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    def _add_key_value_table(self, doc: Document, element: KeyValueTable) -> None:
        if element.title:
            p = doc.add_paragraph(element.title)
            p.runs[0].bold = True
            p.paragraph_format.keep_with_next = True
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        for key, value in element.rows:
            cells = table.add_row().cells
            cells[0].text = key
            cells[1].text = value
            cells[0].width = Cm(5.0)
            cells[1].width = Cm(10.0)
            if cells[0].paragraphs[0].runs:
                cells[0].paragraphs[0].runs[0].bold = True
        self._style_table(table, header=False)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    def _style_table(self, table, header: bool) -> None:
        table.style = 'Table Grid'
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(8)
                if header and r_idx == 0:
                    self._shade_cell(cell, 'D9E2F3')
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    def _add_footer(self, doc: Document, model: DocumentModel) -> None:
        for section in doc.sections:
            footer = section.footer
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.text = f'{model.generated_by} {model.version}'
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(7)

    @staticmethod
    def _repeat_table_header(row) -> None:
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = OxmlElement('w:tblHeader')
        tbl_header.set(qn('w:val'), 'true')
        tr_pr.append(tbl_header)

    @staticmethod
    def _shade_cell(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            tc_pr.append(shd)
        shd.set(qn('w:fill'), fill)

    @staticmethod
    def _set_cell_margins(cell, **kwargs) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in('w:tcMar')
        if tc_mar is None:
            tc_mar = OxmlElement('w:tcMar')
            tc_pr.append(tc_mar)
        for margin in ['top', 'start', 'bottom', 'end']:
            if margin in kwargs:
                node = tc_mar.find(qn(f'w:{margin}'))
                if node is None:
                    node = OxmlElement(f'w:{margin}')
                    tc_mar.append(node)
                node.set(qn('w:w'), str(kwargs.get(margin)))
                node.set(qn('w:type'), 'dxa')
